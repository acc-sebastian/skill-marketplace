#!/usr/bin/env python3
"""
extract_nda_text.py — Extract NDA text with stable paragraph markers.

Usage:
    python extract_nda_text.py <file.docx|file.pdf|file.txt>

Emits UTF-8 text to stdout. Every non-empty paragraph is prefixed with a stable
marker [¶NNN] used for citations in the review report. For PDFs without a text
layer (scans), prints a SCAN-DETECTED notice and exits 2 so the caller falls back
to page rasterization / vision reading.

Dependencies: python-docx (docx), pdftotext (poppler-utils) — both present in the
claude.ai sandbox. Installs nothing by itself.
"""

import re
import subprocess
import sys
from pathlib import Path

MARKER = "[¶{:03d}]"


def emit(paragraphs):
    n = 0
    for p in paragraphs:
        p = p.strip()
        if not p:
            continue
        n += 1
        print(f"{MARKER.format(n)} {p}")
    return n


def extract_docx(path: Path):
    try:
        from docx import Document  # python-docx
    except ImportError:
        sys.exit("python-docx not available — install with: "
                 "pip install python-docx --break-system-packages")
    doc = Document(str(path))
    paras = []
    for p in doc.paragraphs:
        if p.text.strip():
            paras.append(p.text)
    # Tables (signature blocks, party tables) — flatten row-wise
    for t_idx, table in enumerate(doc.tables, 1):
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paras.append(f"(Table {t_idx}) " + " | ".join(dict.fromkeys(cells)))
    return paras


def extract_pdf(path: Path):
    # Text layer present?
    fonts = subprocess.run(["pdffonts", str(path)],
                           capture_output=True, text=True)
    data_lines = [ln for ln in fonts.stdout.splitlines()[2:] if ln.strip()]
    if not data_lines:
        print("SCAN-DETECTED: no text layer. Rasterize pages and read visually "
              "(see pdf-reading skill).", file=sys.stderr)
        sys.exit(2)
    out = subprocess.run(["pdftotext", "-layout", "-enc", "UTF-8",
                          str(path), "-"],
                         capture_output=True, text=True, check=True)
    # Merge hard-wrapped lines into paragraphs: blank line = paragraph break;
    # a line ending in sentence punctuation or a clause number pattern also breaks.
    paras, buf = [], []
    for line in out.stdout.splitlines():
        stripped = line.strip()
        if not stripped:
            if buf:
                paras.append(" ".join(buf))
                buf = []
            continue
        buf.append(re.sub(r"\s{2,}", " ", stripped))
        if re.search(r"[.;:]\s*$", stripped):
            paras.append(" ".join(buf))
            buf = []
    if buf:
        paras.append(" ".join(buf))
    return paras


def extract_txt(path: Path):
    text = path.read_text(encoding="utf-8", errors="replace")
    return re.split(r"\n\s*\n", text)


def main():
    if len(sys.argv) != 2:
        sys.exit(__doc__)
    path = Path(sys.argv[1])
    if not path.exists():
        sys.exit(f"File not found: {path}")
    suffix = path.suffix.lower()
    if suffix == ".docx":
        paras = extract_docx(path)
    elif suffix == ".pdf":
        paras = extract_pdf(path)
    elif suffix in (".txt", ".md"):
        paras = extract_txt(path)
    else:
        sys.exit(f"Unsupported file type: {suffix} (use docx, pdf, or txt)")
    count = emit(paras)
    print(f"\n--- {count} paragraphs extracted from {path.name} ---",
          file=sys.stderr)


if __name__ == "__main__":
    main()
